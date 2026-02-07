# -*- coding: utf-8 -*-
"""
Local HTTP server that proxies Google Cloud Text-to-Speech for the browser viewer.

Why:
- The viewer is a static file (file://), so it can't securely hold service credentials.
- We need word timepoints; Google supports SSML <mark> timepointing.

Usage:
  export GOOGLE_APPLICATION_CREDENTIALS="/path/to/service_account.json"
  python -m src.tts_server

Endpoint:
  POST /tts
  JSON body:
    {
      "ssml": "<speak>...</speak>",
      "language_code": "ar-XA",
      "gender": "MALE",
      "speaking_rate": 1.0,
      "voice_name": null
    }
  Response:
    { "audio_b64": "...", "timepoints": [ {"mark":"w0","time":0.12}, ... ], "voice": "..." }
"""

from __future__ import annotations

import base64
import json
import os
import re
import unicodedata
import datetime
from docx import Document
from http.server import BaseHTTPRequestHandler, HTTPServer
from typing import Any, Dict, Optional, Tuple
from rapidfuzz.distance import Levenshtein
try:
    from openai import OpenAI
except ImportError:
    OpenAI = None
from pathlib import Path
import threading
import threading
from src.config import AUDIO_DIR, AUDIO_MANIFEST, DOC_ARCHIVES_DIR, ALIGNMENT_JSON
from src.services.alignment_service import AlignmentService

alignment_service = AlignmentService()

MANIFEST_LOCK = threading.Lock()


def _json_response(handler: BaseHTTPRequestHandler, code: int, obj: Dict[str, Any]) -> None:
    raw = json.dumps(obj, ensure_ascii=False).encode("utf-8")
    handler.send_response(code)
    handler.send_header("Content-Type", "application/json; charset=utf-8")
    handler.send_header("Content-Length", str(len(raw)))
    handler.send_header("Access-Control-Allow-Origin", "*")
    handler.end_headers()
    handler.wfile.write(raw)


def _read_json(handler: BaseHTTPRequestHandler) -> Tuple[Optional[Dict[str, Any]], Optional[str]]:
    try:
        ln = int(handler.headers.get("Content-Length", "0") or "0")
    except Exception:
        ln = 0
    if ln <= 0:
        return None, "empty body"
    try:
        raw = handler.rfile.read(ln)
        obj = json.loads(raw.decode("utf-8"))
        if not isinstance(obj, dict):
            return None, "body is not a JSON object"
        return obj, None
    except Exception as e:
        return None, f"invalid json: {e}"


_tts_client = None
_voices_cache: Dict[str, Any] = {}
_openai_client = None

def _get_openai_client():
    global _openai_client
    if _openai_client:
        return _openai_client
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return None
    if OpenAI is None:
        print("[TTS Server] OpenAI module not available.")
        return None
    _openai_client = OpenAI(api_key=api_key)
    return _openai_client

def normalize_arabic(text):
    """
    Normalizes Arabic text by keeping ONLY standard Arabic letters.
    Removes EVERYTHING else (numbers, punctuation, symbols, whitespace, harakat, tatweel, etc.).
    Also maps common Persian/Urdu variants to standard Arabic to prevent false mismatches.
    """
    # 1. Map Common Variants (Persian/Urdu -> Arabic)
    text = text.replace("ک", "ك")  # Keheh -> Kaf
    text = text.replace("ی", "ي")  # Persian Yeh -> Arabic Yeh
    text = text.replace("ى", "ي")  # Alef Maksura -> Yeh (Normalize to Yeh for consistency if needed, or keep distinct if strict)
    # Keeping Alef Maksura as distinct usually, checking valid_chars below:
    # valid_chars includes 'ى' (Alef Maksura) and 'ي' (Yeh).
    # So we should probably convert Persian Yeh (ی - 06CC) to Arabic Yeh (ي - 064A)
    # And Keheh (ک - 06A9) to Kaf (ك - 0643)

    # Explicit list of valid Arabic letters as requested
    # Hamza, Madda, Alifs, Hamzas, Beh...Yeh
    valid_chars = "ءآأؤإئابةتثجحخدذرزسشصضطظعغفقكلمنهوىي"
    
    # Create regex to replace anything NOT in valid_chars with empty string
    # Matches any character that is not in the allowed list
    pattern = f"[^{re.escape(valid_chars)}]"
    text = re.sub(pattern, '', text)
    
    return text

def log_to_word(text):
    filename = "test_wordu.docx"
    try:
        if os.path.exists(filename):
            doc = Document(filename)
        else:
            doc = Document()
        doc.add_paragraph("--- New Chunk ---")
        doc.add_paragraph(text)
        doc.save(filename)
    except Exception as e:
        print(f"[TTS Server] Failed to log to word: {e}")

def log_fallback_to_word(segments):
    """
    segments: list of (final_word, rejected_word, is_reverted)
    """
    filename = "test_wordu.docx"
    try:
        if os.path.exists(filename):
            doc = Document(filename)
        else:
            doc = Document()
        
        doc.add_paragraph("--- FALLBACK RESULT (Bold = Reverted) ---")
        p = doc.add_paragraph()
        # Set RTL for better visibility of Arabic
        p.style.font.rtl = True
        
        for text, _, is_bold in segments:
            run = p.add_run(text + " ")
            if is_bold:
                run.bold = True
                run.underline = True
                
        doc.save(filename)
    except Exception as e:
        print(f"[TTS Server] Failed to log fallback to word: {e}")

def log_fallback_to_html(original_text, vocalized_text, norm_orig_list, norm_voc_list, segments, attempt_info="", filename="test_output.html", page_name=None):
    """
    segments: list of (final_word, rejected_word, is_reverted)
    """
    try:
        if not os.path.exists(filename):
            with open(filename, "w", encoding="utf-8") as f:
                f.write("""
<!DOCTYPE html>
<html lang="ar" dir="rtl">
<head>
<meta charset="UTF-8">
<style>
body { font-family: 'Traditional Arabic', 'Amiri', 'Arial', sans-serif; font-size: 20px; direction: rtl; padding: 20px; background: #f9f9f9; }
.section { margin-bottom: 20px; padding: 15px; background: #fff; border: 1px solid #ddd; border-radius: 5px; }
.section h2 { font-size: 18px; color: #333; border-bottom: 1px solid #eee; padding-bottom: 5px; margin-top: 0; }
.reverted { color: red; font-weight: bold; text-decoration: underline; }
.text-content { line-height: 1.8; }
.norm-token { display: inline-block; padding: 2px 5px; margin: 2px; background: #f0f0f0; border-radius: 3px; font-size: 16px; color: #555; }
.norm-empty { background: #ffebeb; color: #999; font-style: italic; }
.meta-info { background: #e3f2fd; color: #0d47a1; padding: 10px; border-radius: 5px; margin-bottom: 15px; font-weight: bold; }
</style>
</head>
<body>
""")

        with open(filename, "a", encoding="utf-8") as f:
            f.write('<div class="chunk-container">')
            
            if page_name:
                f.write(f'<h1 style="background: #e0f7fa; padding: 10px; border-radius: 5px; color: #006064;">Page: {page_name}</h1>')

            if attempt_info:
                f.write(f'<div class="meta-info">{attempt_info}</div>')

            f.write(f'<h2>TTS Log - {attempt_info or "Event"}</h2>')
            
            # 1. Original
            f.write('<div class="section"><h2>Original Text</h2><div class="text-content">')
            f.write(original_text)
            f.write('</div></div>')
            
            # 2. Vocalized (Raw AI Output)
            f.write('<div class="section"><h2>Vocalized Text (AI Raw)</h2><div class="text-content">')
            f.write(vocalized_text)
            f.write('</div></div>')

            # 3. Normalized Comparison (Strict Arabic) -- REMOVED PER USER REQUEST
            # f.write('<div class="section"><h2>Normalized Alignment (Strict Arabic Only)</h2><div class="text-content">')
            # f.write('<strong>Original Skeletons:</strong><br>')
            # for n in norm_orig_list:
            #     cls = "norm-token" + (" norm-empty" if not n else "")
            #     disp = n if n else "[EMPTY]"
            #     f.write(f'<span class="{cls}">{disp}</span> ')
            # f.write('<br><br><strong>Vocalized Skeletons:</strong><br>')
            # for n in norm_voc_list:
            #     cls = "norm-token" + (" norm-empty" if not n else "")
            #     disp = n if n else "[EMPTY]"
            #     f.write(f'<span class="{cls}">{disp}</span> ')
            # f.write('</div></div>')

            # 4. Result (Diff) -- REMOVED PER USER REQUEST
            # f.write('<div class="section"><h2>Result (Diff / Hybrid)</h2><div class="text-content">')
            # for text, _, is_reverted in segments:
            #     if is_reverted:
            #         f.write(f'<span class="reverted">{text}</span> ')
            #     else:
            #         f.write(f'{text} ')
            # f.write('</div>')
            
            # Details if available
            # Details if available
            # Details if available
            reversions = []
            for s in segments:
                if s[2]: # is_reverted
                    # Check if the reverted word has any actual letters.
                    # If it's just punctuation (like "(1),"), ignore it.
                    # We can use normalize_arabic to see if anything remains as a skeleton.
                    skeleton = normalize_arabic(s[0])
                    if skeleton.strip():
                         reversions.append(s)

            f.write('<hr><div style="font-size:16px; color:#666;"><strong>Details (Reversions):</strong><br>')
            if reversions:
                for final, rejected, _ in reversions:
                     f.write(f'<div>Rejected: <span style="text-decoration: line-through;">{rejected}</span> &rarr; Kept: <span class="reverted">{final}</span></div>')
            else:
                f.write('<div><em>Hiçbir değişen kelime yok (Tam eşleşme veya AI kusursuz).</em></div>')
            f.write('</div>')
            
            f.write('</div></div><hr>\n')
            
    except Exception as e:
        print(f"[TTS Server] Failed to log fallback to html: {e}")

def log_google_chunks_to_html(chunks, filename="test_output.html"):
    """
    Logs the split chunks sent to Google TTS to the HTML file.
    """
    try:
        # ensuring we have a valid html header if new file or appending safely
        if not os.path.exists(filename):
             with open(filename, "w", encoding="utf-8") as f:
                 f.write('<!DOCTYPE html><html lang="ar" dir="rtl"><head><meta charset="UTF-8"></head><body>')

        with open(filename, "a", encoding="utf-8") as f:
            f.write('<div class="section" style="background:#eef; border-color:#ccd;"><h2>Google TTS Payloads (Split / Unmerged)</h2><div class="text-content">')
            f.write('<p style="color:#666; font-size:16px;">These are the exact chunks sent to Google TTS (excluding SSML tags for clarity), appearing separately:</p>')
            
            for i, chunk in enumerate(chunks):
                f.write(f'<div style="margin-bottom:10px; padding:10px; background:#fff; border:1px solid #ddd; border-left:4px solid #4CAF50;">')
                f.write(f'<strong>Part {i+1}:</strong><br>')
                f.write(chunk)
                f.write('</div>')
                
            f.write('</div></div>')
    except Exception as e:
        print(f"[TTS Server] Failed to log Google chunks: {e}")

def _count_stats(text: str) -> Tuple[int, int]:
    """
    Returns (letter_count, diacritic_count)
    Strict Arabic Only:
    - Letters: Hamza, Alif, Beh...Yeh and their forms.
    - Diacritics: Fatha, Damma, Kasra, Sukun, Shadda, Tanween, etc.
    """
    valid_letters = "ءآأؤإئابةتثجحخدذرزسشصضطظعغفقكلمنهوىي"
    tashkeel_pattern = r'[\u064B-\u0652\u0670]'

    letter_count = 0
    diacritic_count = 0
    
    for char in text:
        if char in valid_letters:
            letter_count += 1
        elif re.match(tashkeel_pattern, char):
            diacritic_count += 1
            
    return letter_count, diacritic_count

def vocalize_chunk_with_retry(text_chunk: str, log_file_path: str = "test_output.html", page_name: str = None) -> str:
    """
    Vocalizes the text using OpenAI with retries and verification.
    Returns the vocalized text if successful, raises Exception if validation fails.
    """
    client = _get_openai_client()
    if not client:
        print("[TTS Server] OpenAI unavailable (key or module missing). Using original text.")
        return text_chunk
        
    model_name = "gpt-5.2" 
    max_retries = 3
    
    # Normalize original once
    norm_original = normalize_arabic(text_chunk)
    
    # We will track the last vocalized text for logging if needed
    vocalized_text = ""
    
    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": (
                        "You are an expert Arabic linguist. "
                        "Your task is to add full diacritics (Tashkeel) AND correct/add proper punctuation to the following Arabic text. "
                        "Do NOT change any words or their order. "
                        "Return ONLY the fully vocalized and punctuated text."
                    )},
                    {"role": "user", "content": text_chunk}
                ],
                temperature=0
            )
            vocalized_text = response.choices[0].message.content
            if not vocalized_text:
                continue

            # Log for debugging
            log_to_word(vocalized_text)
            
            # 1. Diacritic Density Check (PRIORITY: Retry if poor vocalization)
            l_count, d_count = _count_stats(vocalized_text)
            # Policy: if Letters > 2 * Diacritics, it's likely under-vocalized.
            if l_count > 2 * d_count:
                print(f"[TTS Server] Low diacritic count attempt {attempt+1}/{max_retries}")
                print(f"  Letters: {l_count}, Diacritics: {d_count}")
                
                # Log the FAILED attempt
                try:
                    # Dummy empty segments for log
                    log_fallback_to_html(
                        text_chunk, 
                        vocalized_text, 
                        [normalize_arabic(w) for w in text_chunk.split()], 
                        [normalize_arabic(w) for w in vocalized_text.split()], 
                        [], 
                        [normalize_arabic(w) for w in vocalized_text.split()], 
                        [], 
                        attempt_info=f"Hareke Kontrolü: {attempt+1}. Deneme BAŞARISIZ (Yetersiz Hareke: L={l_count}, D={d_count})",
                        filename=log_file_path,
                        page_name=page_name
                    )
                except Exception: pass

                # Retry to get better vocalization
                continue

            # 2. Structural Check (Mismatch)
            # Logic: If vocalization is GOOD (passed above), but words changed -> Drop to Fallback (do not retry)
            norm_vocalized = normalize_arabic(vocalized_text)
            
            if norm_original != norm_vocalized:
                print(f"[TTS Server] Mismatch detected after passing diacritic check.")
                print(f"  Orig: {norm_original}")
                print(f"  Recv: {norm_vocalized}")
                
                # Log the MISMATCH attempt before breaking
                try:
                    # Quick diff logic for log
                    ws_original = text_chunk.split()
                    ws_vocalized = vocalized_text.split()
                    norm_w_orig = [normalize_arabic(w) for w in ws_original]
                    norm_w_voc = [normalize_arabic(w) for w in ws_vocalized]
                    opcodes = Levenshtein.opcodes(norm_w_orig, norm_w_voc)
                    
                    # We can reuse the main fallback log logic but just for this attempt
                    # We need to construct segments.
                    # Since we are breaking to fallback anyway, we can just let the final fallback log handle the "diff" view.
                    # BUT user wants to see "old trials". Since this BREAKS, there are no more trials.
                    # The "Final Fallback" log below will cover this event.
                    # So we only explicit log here if we were RETRYING. But we are NOT retrying.
                    # We will log it as "Attempt X Mismatch - Going to Fallback"
                    
                    # Let's log it to be safe and explicit as requested "her seferinde"
                    fail_segments = [] 
                    for tag, i1, i2, j1, j2 in opcodes:
                        if tag == 'equal':
                             for k in range(i2-i1): fail_segments.append((ws_vocalized[j1+k], None, False))
                        elif tag == 'replace':
                             for k in range(i2-i1):
                                w_rej = ws_vocalized[j1+k] if k < (j2-j1) else None
                                fail_segments.append((ws_original[i1+k], w_rej, True))
                        elif tag == 'delete':
                             for k in range(i2-i1): fail_segments.append((ws_original[i1+k], "[DEL]", True))
                    
                    log_fallback_to_html(
                        text_chunk, 
                        vocalized_text, 
                        norm_w_orig, 
                        norm_w_voc, 
                        fail_segments, 
                        # fail_segments duplicated arg was in original code? No, simple diff.
                        attempt_info=f"Hareke Kontrolü: BAŞARILI (Ancak {attempt+1}. Denemede Yapısal Uyumsuzluk - Durduruldu)",
                        filename=log_file_path,
                        page_name=page_name
                    )
                except Exception: pass

                # Do NOT retry for mismatch. We have good diacritics, just some wrong letters.
                # Break to let fallback logic restore original words where needed.
                break 
            
            # If all checks pass: SUCCESS
            # Log success to HTML for visibility
            try:
                # Prepare success log segments (all equal)
                ws_original = text_chunk.split()
                ws_vocalized = vocalized_text.split()
                norm_w_orig = [normalize_arabic(w) for w in ws_original]
                norm_w_voc = [normalize_arabic(w) for w in ws_vocalized]
                
                # Since we passed check 1, norms are equal, so Levenshtein should be all 'equal'
                # But to be safe and lazy, we just re-run Levenshtein or assume identity
                # Re-running Levenshtein logic to get segments format:
                opcodes = Levenshtein.opcodes(norm_w_orig, norm_w_voc)
                success_segments = []
                for tag, i1, i2, j1, j2 in opcodes:
                    if tag == 'equal':
                        for k in range(i2-i1):
                            success_segments.append((ws_vocalized[j1+k], None, False))
                    # Mismatches shouldn't really happen here if check passed, but handle gracefully if they do
                    else:
                         for k in range(j2-j1):
                            success_segments.append((ws_vocalized[j1+k], None, False))

                log_fallback_to_html(
                    text_chunk, 
                    vocalized_text, 
                    norm_w_orig, 
                    norm_w_voc, 
                    success_segments, 
                    success_segments, 
                    attempt_info=f"Hareke Kontrolü: {attempt+1}. Denemede BAŞARILI",
                    filename=log_file_path,
                    page_name=page_name
                )
            except Exception as log_err:
                print(f"[TTS Server] Success Log Error: {log_err}")

            return vocalized_text

        except Exception as e:
            print(f"[TTS Server] OpenAI Error: {e}")
            
    # If we reached here, all retries failed (or broken loop).
    # FALLBACK STRATEGY:
    # Compare word-by-word. usage original (unvocalized) word if mismatch.
    print(f"[TTS Server] Failed check or mismatch. Attempting word-level fallback.")
    
    # We use the result from the LAST attempt (vocalized_text) if available.
    if not vocalized_text:
         # If OpenAI completely failed to return text, return pure original
         print("[TTS Server] No response from OpenAI. Returning original text.")
         return text_chunk

    ws_original = text_chunk.split()
    ws_vocalized = vocalized_text.split()
    
    # Global Alignment using Levenshtein
    # 1. Normalize both lists for alignment
    norm_w_orig = [normalize_arabic(w) for w in ws_original]
    norm_w_voc = [normalize_arabic(w) for w in ws_vocalized]
    
    # 2. Get Opcodes
    # opcodes: list of (tag, i1, i2, j1, j2)
    # i = src (original), j = dest (vocalized)
    # We want to conform to 'original' as the source truth for insertions/deletions/replacements
    
    opcodes = Levenshtein.opcodes(norm_w_orig, norm_w_voc)
    
    final_segments = [] # List of (final_word, rejected_word, is_reverted)
    
    for tag, i1, i2, j1, j2 in opcodes:
        if tag == 'equal':
            # Perfect match (normalized)
            for k in range(i2-i1):
                idx_voc = j1 + k
                # Keep the vocalized version
                final_segments.append((ws_vocalized[idx_voc], None, False))
                
        elif tag == 'replace':
            # Mismatch. Revert to original.
            for k in range(i2-i1):
                w_orig = ws_original[i1+k]
                # Try to get corresponding vocalized word for log if exists
                # mapping might be N:M, simplistic 1:1 fallback mapping for log
                w_rej = None
                if k < (j2-j1):
                    w_rej = ws_vocalized[j1+k]
                
                print(f"  [Fallback-Align] Replace: Orig='{w_orig}' vs Voc='{w_rej}'. Reverting.")
                final_segments.append((w_orig, w_rej, True))
                
        elif tag == 'delete':
            # Word in original NOT in vocalized (AI deleted it).
            # Restore it.
            for k in range(i2-i1):
                w_orig = ws_original[i1+k]
                print(f"  [Fallback-Align] Restoring deleted word: '{w_orig}'")
                final_segments.append((w_orig, "[DELETED]", True))
                
        elif tag == 'insert':
            # Word in vocalized NOT in original (AI hallucination/insertion).
            # Ignore it.
            for k in range(j2-j1):
                w_ins = ws_vocalized[j1+k]
                print(f"  [Fallback-Align] Ignoring inserted word: '{w_ins}'")
                # Do not append to final_segments
                
    # Reconstruct text
    final_words = [s[0] for s in final_segments]
    hybrid_text = " ".join(final_words)

    # Log logs
    log_fallback_to_word(final_segments)
    
    # Calculate recovered count (excluding punctuation-only restores)
    reverted_count = 0
    for s in final_segments:
        if s[2]: # passed fallback/revert
             if normalize_arabic(s[0]).strip():
                 reverted_count += 1
    
    log_fallback_to_html(
        text_chunk, 
        vocalized_text, 
        norm_w_orig, 
        norm_w_voc, 
        final_segments,
        attempt_info=f"{reverted_count} kelime orijinalden kurtarıldı",
        filename=log_file_path,
        page_name=page_name
    )
    
    return hybrid_text


def _get_client():
    global _tts_client
    if _tts_client is not None:
        return _tts_client
    try:
        from google.cloud import texttospeech_v1beta1 as texttospeech  # type: ignore
    except Exception as e:
        raise RuntimeError(
            "google-cloud-texttospeech is not installed. Install it and retry. "
            f"Import error: {e}"
        )
    _tts_client = texttospeech.TextToSpeechClient()
    return _tts_client


def _pick_voice(language_code: str, gender: str, voice_name: Optional[str] = None) -> Tuple[str, Any]:
    """
    Pick a voice name for the requested language/gender. Returns (voice_name, voice_params_obj).
    """
    from google.cloud import texttospeech_v1beta1 as texttospeech  # type: ignore

    # Explicit voice wins
    if voice_name:
        vp = texttospeech.VoiceSelectionParams(
            language_code=language_code,
            name=voice_name,
        )
        return voice_name, vp

    # Cache per language
    cache_key = f"{language_code}"
    if cache_key not in _voices_cache:
        client = _get_client()
        try:
            resp = client.list_voices(language_code=language_code)
            _voices_cache[cache_key] = resp.voices or []
        except Exception:
            _voices_cache[cache_key] = []

    voices = _voices_cache.get(cache_key) or []
    want = str(gender or "").upper().strip()
    want_gender = {
        "MALE": texttospeech.SsmlVoiceGender.MALE,
        "FEMALE": texttospeech.SsmlVoiceGender.FEMALE,
        "NEUTRAL": texttospeech.SsmlVoiceGender.NEUTRAL,
    }.get(want, texttospeech.SsmlVoiceGender.MALE)

    chosen = None
    for v in voices:
        try:
            if v and v.ssml_gender == want_gender:
                chosen = v
                break
        except Exception:
            continue

    if chosen is None and voices:
        chosen = voices[0]

    if chosen is not None:
        vp = texttospeech.VoiceSelectionParams(
            language_code=language_code,
            name=getattr(chosen, "name", None) or None,
            ssml_gender=want_gender,
        )
        return getattr(chosen, "name", None) or "", vp

    # Fallback without name (let API choose)
    vp = texttospeech.VoiceSelectionParams(language_code=language_code, ssml_gender=want_gender)
    return "", vp


def split_into_three_by_sentences(text: str) -> list[str]:
    """
    Splits text into 3 roughly equal parts by WORD count, but ONLY breaking at sentence boundaries.
    Respects punctuation: . ? ! ; : ؟ ؛ and newlines.
    """
    if not text:
        return []
    
    # 1. Split into sentences (keep punctuation attached)
    # Regex captures delimiters in a group, so we get [text, delim, text, delim...]
    pattern = r'([.!?;:؟؛\n]+)'
    parts = re.split(pattern, text)
    
    sentences = []
    current = ""
    for p in parts:
        # If p matches the delimiter pattern, append it to current and push to sentences
        if p and re.match(pattern, p):
             current += p
             sentences.append(current.strip())
             current = ""
        else:
             current += p
    
    if current.strip():
        sentences.append(current.strip())
        
    sentences = [s for s in sentences if s]
    
    if not sentences:
        return []
        
    # If very few sentences, can't "split into 3" efficiently, so just return what we have
    if len(sentences) < 3:
        return sentences

    # 2. Calculate word counts to find optimal split points
    counts = [len(s.split()) for s in sentences]
    total_words = sum(counts)
    if total_words == 0:
        return sentences
        
    target1 = total_words / 3.0
    target2 = 2.0 * total_words / 3.0
    
    # Precompute cumulative word counts
    # cumsum[i] = sum(sentences[0]...sentences[i])
    cumsum = []
    c = 0
    for x in counts:
        c += x
        cumsum.append(c)
        
    # 3. Find Cut 1 (end of Part 1)
    # We want index i such that cumsum[i] is closest to target1.
    # i can range from 0 to N-2 (leaving at least 1 for rest)
    best_i = 0
    min_d = float('inf')
    
    # Heuristic: try to ensure each part has at least 1 sentence if possible
    search_end_1 = max(1, len(sentences) - 2) 
    
    for i in range(search_end_1 + 1):
        d = abs(cumsum[i] - target1)
        if d < min_d:
            min_d = d
            best_i = i
            
    # Cut 1 is AT best_i (inclusive). So Part 1 is sentences[:best_i+1].
    
    # 4. Find Cut 2 (end of Part 2)
    # i ranges from best_i+1 to N-2
    best_j = best_i + 1
    min_d = float('inf')
    
    search_start_2 = best_i + 1
    search_end_2 = max(search_start_2, len(sentences) - 1)
    
    for j in range(search_start_2, search_end_2):
        d = abs(cumsum[j] - target2)
        if d < min_d:
            min_d = d
            best_j = j
            
    # Cut 2 is AT best_j (inclusive). So Part 2 is sentences[best_i+1 : best_j+1].
    # Part 3 is sentences[best_j+1 :]
    
    p1 = " ".join(sentences[:best_i+1])
    p2 = " ".join(sentences[best_i+1 : best_j+1])
    p3 = " ".join(sentences[best_j+1 :])
    
    return [p for p in [p1, p2, p3] if p]

class Handler(BaseHTTPRequestHandler):
    def log_message(self, fmt: str, *args) -> None:
        # keep quiet
        return

    def do_OPTIONS(self) -> None:
        self.send_response(204)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()

    def do_POST(self) -> None:
        # --- NEW: Update Line Route ---
        if self.path.rstrip("/") == "/update_line":
            try:
                obj, err = _read_json(self)
                if err:
                    _json_response(self, 400, {"error": err})
                    return
                
                line_no = obj.get("line_no")
                new_text = obj.get("new_text")
                
                if not isinstance(line_no, int) or not isinstance(new_text, str):
                    _json_response(self, 400, {"error": "Invalid params"})
                    return

                success = alignment_service.update_line(line_no, new_text)
                
                if success:
                    _json_response(self, 200, {"ok": True})
                else:
                    _json_response(self, 404, {"error": "Line not found or save failed"})
            except Exception as e:
                print(f"[TTS Server] Update lines error: {e}")
                _json_response(self, 500, {"error": str(e)})
            return

        if self.path.rstrip("/") != "/tts":
            _json_response(self, 404, {"error": "not_found"})
            return

        obj, err = _read_json(self)
        if err:
            _json_response(self, 400, {"error": err})
            return

        ssml = (obj.get("ssml") or "").strip()
        tokens = obj.get("tokens")
        if tokens is not None and not isinstance(tokens, list):
            tokens = None

        if not ssml and not tokens:
            _json_response(self, 400, {"error": "ssml or tokens is required"})
            return

        language_code = (obj.get("language_code") or "ar-XA").strip() or "ar-XA"
        gender = (obj.get("gender") or "MALE").strip() or "MALE"
        voice_name = obj.get("voice_name")
        if voice_name is not None and not isinstance(voice_name, str):
            voice_name = None
        speaking_rate = obj.get("speaking_rate", 1.0)
        try:
            speaking_rate = float(speaking_rate)
        except Exception:
            speaking_rate = 1.0
        speaking_rate = max(0.50, min(1.25, speaking_rate))
        chunk_size = obj.get("chunk_size", 450)
        try:
            chunk_size = int(chunk_size)
        except Exception:
            chunk_size = 450
        chunk_size = max(50, min(900, chunk_size))
        token_start = obj.get("token_start", 0)
        try:
            token_start = int(token_start)
        except Exception:
            token_start = 0
        
        action = obj.get("action")
        page_key = obj.get("page_key")
        archive_path_name = obj.get("archive_path")
        reset_log = obj.get("reset_log") # boolean or null

        # --- Lazy Archive Cache Lookup ---
        # If not "batch_save", but we have archive_path and page_key, check if valid audio exists on disk.
        if action != "batch_save" and page_key and archive_path_name:
            print(f"DEBUG: Lazy Lookup Limit Check. Archive: '{archive_path_name}', Page: '{page_key}'")
            try:
                # Fuzzy match for directory name to handle trailing spaces/unicode
                # Normalize input to NFC
                norm_input = unicodedata.normalize('NFC', archive_path_name).strip()
                
                target_dir = None
                # Scan all directories
                if DOC_ARCHIVES_DIR.exists():
                    for d in DOC_ARCHIVES_DIR.iterdir():
                        if d.is_dir():
                            # Normalize dir name to NFC
                            norm_d = unicodedata.normalize('NFC', d.name).strip()
                            if norm_d == norm_input:
                                target_dir = d
                                print(f"DEBUG: Match found! '{d.name}' (norm: '{norm_d}') matches input '{archive_path_name}'")
                                break
                
                if not target_dir:
                    print(f"DEBUG: No folder matched for input '{archive_path_name}' (norm: '{norm_input}')")
                    # Fallback to direct path just in case
                    target_dir = DOC_ARCHIVES_DIR / archive_path_name

                # Determine manifest based on nusha_id
                nusha_id = obj.get("nusha_id")
                try: nusha_id = int(nusha_id)
                except: nusha_id = 1
                
                manifest_filename = "audio_manifest.json"
                audio_subdir = ""
                if nusha_id > 1:
                    manifest_filename = f"audio_manifest_n{nusha_id}.json"
                    audio_subdir = f"n{nusha_id}/"

                target_manifest = target_dir / manifest_filename
                print(f"DEBUG: Checking manifest at: {target_manifest}")
                if target_manifest.exists():
                    try:
                        manifest = json.loads(target_manifest.read_text(encoding="utf-8"))
                        if page_key in manifest:
                            print(f"DEBUG: Page key found in manifest.")
                            chunks_data = manifest[page_key]
                            
                            output_chunks = []
                            all_ok = True
                            for c in chunks_data:
                                rel_path = c.get("audio_path")
                                if not rel_path: 
                                    all_ok = False; break
                                
                                # Resolve path relative to archive
                                # audio_path is like "audio/p3_chunk_0.mp3" OR "audio/n2/p3_chunk_0.mp3"
                                # If the path in manifest is relative to archive root, we just join it.
                                full_path = target_dir / rel_path
                                print(f"DEBUG: Checking audio file: {full_path}")
                                
                                if not full_path.exists():
                                    print(f"DEBUG: Audio file MISSING: {full_path}")
                                    all_ok = False
                                    break
                                
                                # Read and encode
                                audio_bytes = full_path.read_bytes()
                                b64 = base64.b64encode(audio_bytes).decode('utf-8')
                                output_chunks.append({
                                    "audio_b64": b64,
                                    "timepoints": c.get("timepoints")
                                })
                            
                            if all_ok and output_chunks:
                                print(f"DEBUG: Lazy lookup SUCCESS. Returning {len(output_chunks)} chunks.")
                                self.send_response(200)
                                self.send_header('Content-type', 'application/json')
                                self.end_headers()
                                self.wfile.write(json.dumps({
                                    "chunks": output_chunks,
                                    "source": "archive_cache"
                                }).encode('utf-8'))
                                return
                        else:
                            print(f"DEBUG: Page key '{page_key}' NOT found in manifest.")
                    except Exception as e:
                        print(f"DEBUG: Error reading manifest: {e}")
                else:
                    print(f"DEBUG: Manifest file does not exist.")
            except Exception as e:
                print(f"DEBUG: Error in lazy lookup: {e}")

        # If it was a "check_only" request and we reached here (lazy lookup failed), return 404.
        if action == "check_only":
             _json_response(self, 404, {"error": "Audio not prepared (check_only)."})
             return
        # ---------------------------------

        try:
            from google.cloud import texttospeech_v1beta1 as texttospeech  # type: ignore

            client = _get_client()
            chosen_name, voice = _pick_voice(language_code, gender, voice_name=voice_name)

            audio_config = texttospeech.AudioConfig(
                audio_encoding=texttospeech.AudioEncoding.MP3,
                speaking_rate=speaking_rate,
            )

            def _synth_one(ssml_text: str) -> Dict[str, Any]:
                synthesis_input = texttospeech.SynthesisInput(ssml=ssml_text)
                req = texttospeech.SynthesizeSpeechRequest(
                    input=synthesis_input,
                    voice=voice,
                    audio_config=audio_config,
                    enable_time_pointing=[texttospeech.SynthesizeSpeechRequest.TimepointType.SSML_MARK],
                )
                resp = client.synthesize_speech(request=req)
                audio_b64 = base64.b64encode(resp.audio_content or b"").decode("ascii")
                tps = []
                for tp in (resp.timepoints or []):
                    try:
                        tps.append({"mark": tp.mark_name, "time": float(tp.time_seconds)})
                    except Exception:
                        continue
                return {"audio_b64": audio_b64, "timepoints": tps}

            # If tokens provided, chunk by N words with NO punctuation insertion.
            if tokens:
                toks = [str(t).strip() for t in tokens if str(t).strip()]
                
                # OpenAI Chunk Size (can be reasonably large, OpenAI handles 4k tokens ~ 2-3k words easily)
                # But we stick to smaller chunks to be safe and responsive.
                openai_chunk_size = 300 
                
                # Prepare output structure
                # We return a list of chunks, or SAVE them if batch_save
                final_chunks_output = []
                
                # Global index tracker for 'w{i}' marks relative to the start of the ENTIRE document (or provided token_start)
                current_global_token_index = token_start
                
                # Pre-calculate audio filename prefix if batch saving
                batch_files_created = []

                # Handle log reset before chunk loop
                if reset_log:
                    lg_p = "test_output.html"
                    if archive_path_name:
                         lg_p = str(DOC_ARCHIVES_DIR / archive_path_name / "test_output.html")
                    try:
                        if os.path.exists(lg_p):
                            os.remove(lg_p)
                    except Exception: pass
                    # Turn off reset_log so we don't delete for next chunk in this request
                    reset_log = False

                for i in range(0, len(toks), openai_chunk_size):
                    part = toks[i : i + openai_chunk_size]
                    
                    # 1. Reconstruct text for OpenAI
                    # We join with spaces. Strict assumption: tokens are words.
                    chunk_text_raw = " ".join(part)
                    
                    # 2. Vocalize with retry
                    # Determine log path
                    log_path = "test_output.html"
                    if archive_path_name:
                         log_path = str(DOC_ARCHIVES_DIR / archive_path_name / "test_output.html")
                    
                    # If reset_log is requested for this batch run (first page), delete now
                    if reset_log:
                        try:
                            # We only do this ONCE. 
                            # But wait, we are in a loop for chunks. We should do it before loop.
                            pass 
                        except: pass
                        
                    try:
                        vocalized_text = vocalize_chunk_with_retry(chunk_text_raw, log_file_path=log_path, page_name=page_key)
                    except Exception as e:
                        _json_response(self, 500, {"error": f"Vocalization failed: {str(e)}"})
                        return
                        
                    # 3. Split into 3 Equal Parts for Google TTS (Limit 5000 bytes)
                    # User requested: roughly 3 equal parts, BUT split at sentence endings (punctuation).
                    sub_chunks = split_into_three_by_sentences(vocalized_text)
                    
                    # Log the SPLIT version (unmerged) as requested
                    log_google_chunks_to_html(sub_chunks, filename=log_path)
                    
                    # We reuse the logic: Iterate through sub_chunks, generate SSML, send to Google.
                    # Safety check: each sentence-based part is likely < 200 words (since OpenAI chunk is 300).
                    # 4000 byte limit is easily respected.
                    
                    voc_idx_counter = 0 # Relative to this 'part' (original tokens)
                    
                    for sc in sub_chunks:
                        # Note: sc is a sentence or group of sentences.
                        # We split by whitespace to re-align with tokens.
                        sc_words = sc.split()
                        if not sc_words: continue
                        
                        # Build SSML for this sub-chunk
                        sc_ssml_fragments = []
                        for t in sc_words:
                            # Determine global index
                            if voc_idx_counter < len(part):
                                real_idx = current_global_token_index + voc_idx_counter
                                mark = f'<mark name="w{real_idx}"/>'
                                voc_idx_counter += 1
                            else:
                                mark = ""
                                
                            escaped_word = _escape_xml(t)
                            sc_ssml_fragments.append(f'{mark}{escaped_word}')
                        
                        # Create final SSML for this sub-chunk
                        sc_ssml_body = " ".join(sc_ssml_fragments)
                        # Use full <speak> tag
                        final_ssml = f"<speak>{sc_ssml_body}</speak>"
                        
                        try:
                            out = _synth_one(final_ssml)
                            if action == "batch_save" and page_key:
                                # Determine Target Dir
                                if archive_path_name:
                                    # Safe safety check?
                                    # Assuming archive_path_name is just the folder name e.g. "2026..."
                                    target_audio_dir = DOC_ARCHIVES_DIR / archive_path_name / "audio"
                                    
                                    # Nusha isolation
                                    nusha_id = obj.get("nusha_id")
                                    try: nusha_id = int(nusha_id)
                                    except: nusha_id = 1
                                    
                                    if nusha_id > 1:
                                        target_audio_dir = target_audio_dir / f"n{nusha_id}"
                                else:
                                    target_audio_dir = AUDIO_DIR
                                
                                target_audio_dir.mkdir(parents=True, exist_ok=True)
                                
                                # Save to disk
                                filename = f"{page_key}_chunk_{len(batch_files_created)}.mp3"
                                file_path = target_audio_dir / filename
                                
                                # Relative path for browser
                                # standard: audio/filename.mp3
                                # isolated: audio/n2/filename.mp3
                                rel_prefix = "audio/"
                                if nusha_id > 1:
                                    rel_prefix = f"audio/n{nusha_id}/"
                                    
                                with open(file_path, "wb") as f:
                                    f.write(base64.b64decode(out["audio_b64"]))
                                
                                # We strip audio_b64 from response to save bandwidth, just keep path
                                out_saved = {
                                    "audio_path": f"{rel_prefix}{filename}", # Relative for browser
                                    "timepoints": out["timepoints"]
                                }
                                batch_files_created.append(out_saved)
                            else:
                                final_chunks_output.append(out)
                        except Exception as syn_err:
                            print(f"[TTS Server] Google Synth Error: {syn_err}")
                
                # Advance global index by the amount of ORIGINAL tokens we consumed in this chunk (should be len(part))
                # Note: if voc_idx_counter < len(part), we technically missed some markers?
                # But we must keep the global index consistent with the INPUT tokens.
                current_global_token_index += len(part)

            if action == "batch_save" and page_key and batch_files_created:
                # Update Manifest
                with MANIFEST_LOCK:
                    nusha_id = obj.get("nusha_id")
                    try: nusha_id = int(nusha_id)
                    except: nusha_id = 1

                    if archive_path_name:
                        manifest_filename = "audio_manifest.json"
                        if nusha_id > 1:
                            manifest_filename = f"audio_manifest_n{nusha_id}.json"
                        target_manifest = DOC_ARCHIVES_DIR / archive_path_name / manifest_filename
                    else:
                        target_manifest = AUDIO_MANIFEST

                    manifest = {}
                    if target_manifest.exists():
                        try: manifest = json.loads(target_manifest.read_text(encoding="utf-8"))
                        except: pass
                    manifest[page_key] = batch_files_created
                    target_manifest.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
                
                _json_response(self, 200, {"ok": True, "saved_chunks": batch_files_created})
                return

            _json_response(self, 200, {"chunks": final_chunks_output, "voice": chosen_name})
            return

            # Single SSML mode (backward compat)
            out = _synth_one(ssml)
            _json_response(self, 200, {"audio_b64": out["audio_b64"], "timepoints": out["timepoints"], "voice": chosen_name})
        except Exception as e:
            _json_response(self, 500, {"error": str(e)})


def _escape_xml(s: str) -> str:
    return (
        (s or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&apos;")
    )


def serve(host: str = "127.0.0.1", port: int = 8765) -> None:
    httpd = HTTPServer((host, port), Handler)
    httpd.serve_forever()


if __name__ == "__main__":
    host = os.environ.get("PATRON_TTS_HOST", "127.0.0.1")
    try:
        port = int(os.environ.get("PATRON_TTS_PORT", "8765"))
    except Exception:
        port = 8765
    serve(host=host, port=port)
